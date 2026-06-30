from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x40, 0x5A)

code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)

def add_code(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='Code')

def add_path(path):
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(2)

# ── TITLE ───────────────────────────────────────────────────────────
title = doc.add_heading('Titan Travel — Tutorial Coding Fundamental', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'Panduan belajar codebase Next.js 16 + React 19 + TypeScript + Prisma + Tailwind v4',
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════
# BAGIAN 1 — BACKEND
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Bagian 1: Backend', level=1)

# ── 1.1 Prisma Schema & Database ──────────────────────────────────
doc.add_heading('1.1 Prisma Schema & Database', level=2)
doc.add_paragraph(
    'Project ini menggunakan PostgreSQL via Neon dengan Prisma v7 sebagai ORM. '
    'Prisma client di-generate ke generated/prisma/client (bukan @prisma/client). '
    'Adapter yang dipakai adalah PrismaPg.'
)
doc.add_paragraph('Langkah-langkah memahami Prisma di project ini:', style='List Bullet')
doc.add_paragraph('Buka prisma/schema.prisma — definisi semua model (User, Booking, TourPackage, dll.)', style='List Bullet')
doc.add_paragraph('Setiap tabel punya field dengan tipe data Prisma: String, Int, Float, Decimal, Json, enum', style='List Bullet')
doc.add_paragraph('Json dipakai untuk field multi-bahasa: { en: string, id: string, ms: string }', style='List Bullet')
doc.add_paragraph('Relasi: Booking → TourPackage (via packageId), PriceTier → TourPackage (via packageId)', style='List Bullet')
doc.add_paragraph('Enum: Role (USER, MANAGER, ADMIN), BookingStatus (PENDING, CONFIRMED, CANCELLED, COMPLETED), PaymentType (DP, HALF, FULL)', style='List Bullet')
doc.add_paragraph('')

add_path('prisma/schema.prisma')
add_code('''enum Role { USER MANAGER ADMIN }

model TourPackage {
  id             String     @id @default(cuid())
  slug           String     @unique
  title          Json       // { en: "...", id: "...", ms: "..." }
  description    Json
  location       Json
  duration       Json
  capacity       Int
  images         Json       // string[]
  facilityScore  Int        // 1-5
  durationDays   Int
  isPublished    Boolean    @default(false)
  bookings       Booking[]
  priceTiers     PriceTier[]
  createdAt      DateTime   @default(now())
  updatedAt      DateTime   @updatedAt
}''')
doc.add_paragraph('')

doc.add_heading('1.2 Singleton Prisma Client', level=2)
doc.add_paragraph(
    'src/lib/prisma.ts membuat instance PrismaClient global (singleton) '
    'agar tidak membuat koneksi baru setiap hot-reload di development.'
)
add_path('src/lib/prisma.ts')
add_code('''import { PrismaClient } from "generated/prisma/client";
import { PrismaPg } from "@prisma/adapter-pg";
import { Pool } from "pg";

const globalForPrisma = globalThis as unknown as { prisma: PrismaClient };
const pool = new Pool({ connectionString: process.env.DATABASE_URL });
const adapter = new PrismaPg(pool);

export const prisma = globalForPrisma.prisma ?? new PrismaClient({ adapter });
if (process.env.NODE_ENV !== "production") globalForPrisma.prisma = prisma;''')
doc.add_paragraph('')

# ── 1.2 Auth ───────────────────────────────────────────────────────
doc.add_heading('1.3 Auth System (JWT Custom)', level=2)
doc.add_paragraph(
    'Tidak menggunakan NextAuth. Auth memakai jwt buatan sendiri via library jose. '
    'Token disimpan di cookie http-only bernama auth_token. Flow: login → encrypt payload → set cookie → '
    'setiap request protected → getSession() membaca cookie → decrypt → dapat data user.'
)
add_path('src/lib/auth.ts')
add_code('''import { SignJWT, jwtVerify } from "jose";
import { cookies } from "next/headers";

const key = new TextEncoder().encode(process.env.JWT_SECRET);

export async function encrypt(payload: JWTPayload) {
  return await new SignJWT(payload)
    .setProtectedHeader({ alg: "HS256" })
    .setExpirationTime("1d")
    .sign(key);
}

export async function decrypt(token: string) {
  try {
    const { payload } = await jwtVerify(token, key, { algorithms: ["HS256"] });
    return payload;
  } catch { return null; }
}

export async function getSession() {
  const token = (await cookies()).get("auth_token")?.value;
  return token ? decrypt(token) : null;
}''')
doc.add_paragraph('')

# ── 1.3 API Routes ─────────────────────────────────────────────────
doc.add_heading('1.4 API Routes (Backend Entry Point)', level=2)
doc.add_paragraph(
    'Hanya ada 6 endpoint API. Selebihnya logic dikerjakan via Server Actions. '
    'API dipakai untuk operasi yang dipicu dari client eksternal (login, register, upload, translate).'
)
doc.add_paragraph('')
doc.add_paragraph('CRUD sederhana — API Auth:', style='List Bullet')
add_path('src/app/api/auth/login/route.ts')
add_code('''export async function POST(req: Request) {
  const { email, password } = await req.json();
  const user = await prisma.user.findUnique({ where: { email } });
  if (!user || !(await bcrypt.compare(password, user.password))) {
    return Response.json({ error: "Invalid credentials" }, { status: 401 });
  }
  const token = await encrypt({ id: user.id, role: user.role });
  (await cookies()).set("auth_token", token, {
    httpOnly: true, secure: true, maxAge: 86400, path: "/",
  });
  return Response.json({ user: { id: user.id, name: user.name, email: user.email, role: user.role } });
}''')
doc.add_paragraph('')
doc.add_paragraph('Struktur API route di Next.js App Router:', style='List Bullet')
add_code('''src/app/api/auth/login/route.ts    → POST /api/auth/login
src/app/api/auth/logout/route.ts   → POST /api/auth/logout
src/app/api/auth/register/route.ts → POST /api/auth/register
src/app/api/auth/session/route.ts  → GET  /api/auth/session
src/app/api/translate/route.ts     → POST /api/translate
src/app/api/upload/route.ts        → POST /api/upload''')
doc.add_paragraph('')

# ── 1.4 Server Actions ─────────────────────────────────────────────
doc.add_heading('1.5 Server Actions (Business Logic)', level=2)
doc.add_paragraph(
    'Server Actions adalah fungsi async yang jalan di server, bisa dipanggil langsung dari komponen client. '
    'Ini adalah pola utama di project ini — 17 file actions di src/lib/actions/. '
    'Setiap action pakai "use server" di baris pertama.'
)
doc.add_paragraph('')
doc.add_paragraph('Pola CRUD — ambil contoh src/lib/actions/packages.ts:', style='List Bullet')
add_code('''"use server";
import { prisma } from "@/lib/prisma";
import { revalidatePath } from "next/cache";

// CREATE
export async function createPackage(data: PackageInput) {
  const slug = data.title.en.toLowerCase().replace(/\\s+/g, "-");
  const pkg = await prisma.tourPackage.create({
    data: { ...data, slug },
  });
  revalidatePath("/admin/packages");
  return pkg;
}

// READ
export async function getPackages() {
  return prisma.tourPackage.findMany({
    include: { priceTiers: true },
    orderBy: { createdAt: "desc" },
  });
}

// UPDATE
export async function updatePackage(id: string, data: Partial<PackageInput>) {
  return prisma.tourPackage.update({ where: { id }, data });
}

// DELETE
export async function deletePackage(id: string) {
  await prisma.tourPackage.delete({ where: { id } });
  revalidatePath("/admin/packages");
}''')
doc.add_paragraph('')
doc.add_paragraph('Daftar 17 Server Actions:', style='List Bullet')
actions_list = [
    'account.ts — Update profile, change password',
    'bank-accounts.ts — CRUD bank accounts',
    'bookings.ts — Create, update, confirm, cancel booking',
    'dashboard.ts — Get dashboard stats (counts, revenue)',
    'gallery.ts — CRUD gallery images',
    'notifications.ts — Create, mark-read notifications',
    'packages.ts — CRUD tour packages + price tiers',
    'recommendation.ts — Rekomendasi paket via TOPSIS',
    'reports.ts — Generate laporan booking (Excel)',
    'settings.ts — Update site settings (hero, about, contact, services)',
    'stats.ts — Statistik revenue bulanan/tahunan',
    'testimonials.ts — CRUD testimonials',
    'topsis.ts — Run TOPSIS analysis, update criteria',
    'translate.ts — Auto-translate via Google Translate API',
    'upload.ts — File upload (Vercel Blob + local fallback)',
    'users.ts — Manage users (admin/manager)',
    'vehicle-types.ts — CRUD vehicle types',
]
for a in actions_list:
    doc.add_paragraph(a, style='List Bullet')
doc.add_paragraph('')

# ── 1.5 TOPSIS Engine (DETAIL) ────────────────────────────────────
doc.add_heading('1.6 TOPSIS Engine (Detail)', level=2)
doc.add_paragraph(
    'TOPSIS (Technique for Order Preference by Similarity to Ideal Solution) '
    'adalah metode decision support untuk meranking alternatif berdasarkan kedekatan '
    'dengan solusi ideal. Di project ini dipakai untuk merekomendasikan paket wisata '
    'terbaik berdasarkan 4 kriteria.'
)
doc.add_paragraph('')

doc.add_heading('Kriteria TOPSIS', level=3)
doc.add_paragraph('Empat kriteria dengan bobot default:', style='List Bullet')

add_code('''C1 — Harga (Price)      → COST    → Bobot 0.35
C2 — Fasilitas (Facilities) → BENEFIT → Bobot 0.25
C3 — Waktu Keberangkatan    → COST    → Bobot 0.20
C4 — Durasi (Duration)      → COST    → Bobot 0.20''')
doc.add_paragraph('')
doc.add_paragraph(
    'Bobot dan tipe (COST/BENEFIT) disimpan di tabel topsis_criteria, '
    'bisa diubah dari panel admin tanpa修改 kode.'
)
doc.add_paragraph('')

doc.add_heading('Interface & Tipe Data', level=3)
add_path('src/lib/topsis.ts')
add_code('''export interface TopsisAlternative {
  id: string;
  name: string;
  c1_price: number;       // Harga (Cost)
  c2_facilities: number;  // Fasilitas 1-5 (Benefit)
  c3_departure: number;   // 1=Pagi, 2=Siang, 3=Malam (Cost)
  c4_duration: number;    // Durasi dalam hari (Cost)
}

export interface TopsisConfig {
  c1: { weight: number; type: "COST" | "BENEFIT" };
  c2: { weight: number; type: "COST" | "BENEFIT" };
  c3: { weight: number; type: "COST" | "BENEFIT" };
  c4: { weight: number; type: "COST" | "BENEFIT" };
}''')
doc.add_paragraph('')

doc.add_heading('Algoritma — 7 Langkah TOPSIS', level=3)

steps = [
    ('Langkah 1: Normalisasi Matriks (R)',
     'Menghitung pembagi (akar dari jumlah kuadrat) per kriteria. '
     'Rumus: R_ij = x_ij / sqrt(Σ x_ij²) untuk setiap kolom.',
     '''const sums = { c1: 0, c2: 0, c3: 0, c4: 0 };
alternatives.forEach(alt => {
  sums.c1 += Math.pow(alt.c1_price, 2);
  sums.c2 += Math.pow(alt.c2_facilities, 2);
  sums.c3 += Math.pow(alt.c3_departure, 2);
  sums.c4 += Math.pow(alt.c4_duration, 2);
});
const divisors = {
  c1: Math.sqrt(sums.c1),
  c2: Math.sqrt(sums.c2),
  c3: Math.sqrt(sums.c3),
  c4: Math.sqrt(sums.c4),
};'''),

    ('Langkah 2 & 3: Matriks Ternormalisasi Terbobot (Y)',
     'Setiap nilai normalisasi dikalikan bobot kriteria. Y_ij = R_ij * W_j.',
     '''const weightedMatrix = alternatives.map(alt => ({
  y1: (alt.c1_price / divisors.c1) * config.c1.weight,
  y2: (alt.c2_facilities / divisors.c2) * config.c2.weight,
  y3: (alt.c3_departure / divisors.c3) * config.c3.weight,
  y4: (alt.c4_duration / divisors.c4) * config.c4.weight,
}));'''),

    ('Langkah 4: Solusi Ideal Positif (A+) & Negatif (A-)',
     'A+ = nilai terbaik (max untuk BENEFIT, min untuk COST). A- = kebalikannya.',
     '''const aPlus = {
  y1: config.c1.type === "COST"
    ? Math.min(...y1_values) : Math.max(...y1_values),
  y2: config.c2.type === "COST"
    ? Math.min(...y2_values) : Math.max(...y2_values),
  // ... sama untuk y3, y4
};'''),

    ('Langkah 5: Jarak Euclidean (D+ dan D-)',
     'Hitung jarak setiap alternatif ke A+ dan A-. '
     'D_i⁺ = √[ Σ (Y_ij — A⁺_j)² ]',
     '''const dPlus = Math.sqrt(
  Math.pow(m.y1 - aPlus.y1, 2) +
  Math.pow(m.y2 - aPlus.y2, 2) +
  Math.pow(m.y3 - aPlus.y3, 2) +
  Math.pow(m.y4 - aPlus.y4, 2)
);

const dMinus = Math.sqrt(
  Math.pow(m.y1 - aMinus.y1, 2) +
  Math.pow(m.y2 - aMinus.y2, 2) +
  Math.pow(m.y3 - aMinus.y3, 2) +
  Math.pow(m.y4 - aMinus.y4, 2)
);'''),

    ('Langkah 6: Nilai Preferensi (Vi)',
     'Vi = D⁻ / (D⁺ + D⁻). Semakin mendekati 1, semakin baik.',
     'const score = dMinus / (dPlus + dMinus) || 0;'),

    ('Langkah 7: Perankingan',
     'Urutkan Vi dari terbesar ke terkecil. Ranking 1 = rekomendasi terbaik.',
     '''return results
  .sort((a, b) => b.score - a.score)
  .map((res, index) => ({ ...res, ranking: index + 1 }));'''),
]

for title_text, desc, code in steps:
    doc.add_heading(title_text, level=3)
    doc.add_paragraph(desc)
    add_code(code)
    doc.add_paragraph('')

doc.add_heading('Badge Hasil TOPSIS', level=3)
doc.add_paragraph(
    'Fungsi getTopsisBadge() mengelompokkan skor ke 4 kategori untuk ditampilkan di UI:'
)
add_code('''Ranking 1  → "Sangat Direkomendasikan" (hijau)
Skor ≥ 0.75 → "Direkomendasikan" (biru)
Skor ≥ 0.50 → "Cukup Sesuai" (kuning)
Skor < 0.50 → "Pilihan Alternatif" (abu-abu)''')
doc.add_paragraph('')

doc.add_heading('TOPSIS Server Action', level=3)
add_path('src/lib/actions/topsis.ts')
doc.add_paragraph(
    'Mengambil data criteria dari DB, mapping alternatif dari TourPackage, '
    'panggil calculateTopsis(), simpan hasil untuk ditampilkan di halaman admin.'
)
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════
# BAGIAN 2 — FRONTEND
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Bagian 2: Frontend', level=1)

# ── 2.1 Route Structure ───────────────────────────────────────────
doc.add_heading('2.1 Route Structure (App Router)', level=2)
doc.add_paragraph(
    'Next.js 16 App Router: setiap folder = route segment. '
    'page.tsx = yang di-render. layout.tsx = wrapper bersama. '
    'Route group (parentheses) tidak ngaruh ke URL, cuma untuk organisasi.'
)
add_code('''src/app/
├── page.tsx                    → /
├── layout.tsx                  → Root layout (fonts, ThemeProvider, metadata)
├── globals.css                 → Global styles + design tokens
├── login/page.tsx              → /login
├── register/page.tsx           → /register
├── dashboard/                  → /dashboard (user)
├── paket/                      → /paket (catalog)
│   └── [slug]/                 → /paket/:slug (detail)
│       └── booking/            → /paket/:slug/booking
├── (panel)/                    → Route group (tidak ngaruh ke URL)
│   ├── admin/                  → /admin/*
│   └── manager/                → /manager/*
└── api/                        → /api/* (REST endpoints)''')
doc.add_paragraph('')

# ── 2.2 Layout ─────────────────────────────────────────────────────
doc.add_heading('2.2 Layout & Providers', level=2)
doc.add_paragraph(
    'Root layout (src/app/layout.tsx) menyediakan font, metadata, dan provider:'
)
add_code('''// Fonts: Playfair Display (heading), DM Sans (body), Geist (mono)
<html lang="id">
  <ThemeProvider>      // dark/light mode via context
    <LocaleProvider>   // en/id/ms via context
      {children}
      <Toaster />      // notifikasi toast (sonner)
    </LocaleProvider>
  </ThemeProvider>
</html>''')
doc.add_paragraph('')
doc.add_paragraph(
    'Panel layout (src/app/(panel)/layout.tsx) melindungi rute admin/manager: '
    'cek session, redirect ke /login jika belum login, cek role (hanya ADMIN/MANAGER). '
    'Menyediakan Sidebar + PanelHeader + NotificationToast.'
)
doc.add_paragraph('')

# ── 2.3 Globals CSS ────────────────────────────────────────────────
doc.add_heading('2.3 Global CSS & Design Tokens', level=2)
add_path('src/app/globals.css')
doc.add_paragraph(
    'Menggunakan Tailwind v4 dengan @import "tailwindcss". Design tokens '
    'didefinisikan di @theme inline block:'
)
add_code('''@import "tailwindcss";
@import "tw-animate-css";
@import "shadcn/tailwind.css";

@custom-variant dark (&:is(.dark *));

@theme inline {
  --color-primary-500: #E8630F;     // Warm Orange (brand utama)
  --color-primary-600: #D4550A;
  --color-accent-500: #F59E0B;      // Warm Amber
  --color-neutral-50: #F8FAFC;
  // ... scale 50-950 untuk setiap palette
  --color-background-secondary: #F1F5F9;
  --color-foreground-secondary: #64748B;
  --color-card-bg: #FFFFFF;
}

[data-theme="dark"] .dark {
  --color-primary-500: #F97316;     // Lebih terang utk dark mode
  --color-background-secondary: #1E293B;
}''')
doc.add_paragraph('')
doc.add_paragraph('Utility classes kustom:', style='List Bullet')
doc.add_paragraph('.glass — efek transparan dengan backdrop blur', style='List Bullet')
doc.add_paragraph('.gradient-text — gradien pada teks heading', style='List Bullet')
doc.add_paragraph('.section-padding — padding konsisten antar section', style='List Bullet')
doc.add_paragraph('.btn-primary / .btn-secondary / .btn-accent — variasi tombol', style='List Bullet')
doc.add_paragraph('Animasi: fadeInUp, fadeIn, slideInLeft, slideInRight, float, pulse-glow, shimmer', style='List Bullet')
doc.add_paragraph('')

# ── 2.4 Component Structure ────────────────────────────────────────
doc.add_heading('2.4 Component Structure', level=2)
doc.add_paragraph(
    'Tiga kategori komponen: UI primitives (shadcn/ui), flat components (layout), '
    'dan panel components (admin/manager). Total 47 komponen.'
)
doc.add_paragraph('')

doc.add_heading('UI Primitives (shadcn/ui)', level=3)
add_path('src/components/ui/')
doc.add_paragraph(
    'Dibangun di atas base-ui/react. Polanya: gunakan class-variance-authority (cva) '
    'untuk variant props, tailwind-merge + clsx untuk class merging.'
)
add_code('''// Contoh: button.tsx
const buttonVariants = cva(
  "inline-flex items-center justify-center rounded-lg font-medium transition-colors",
  {
    variants: {
      variant: {
        default: "bg-primary-500 text-white hover:bg-primary-600",
        outline: "border border-primary-200 text-primary-700",
        ghost: "text-slate-600 hover:bg-slate-100",
      },
      size: { sm: "h-9 px-3 text-sm", md: "h-10 px-4", lg: "h-12 px-6" },
    },
  }
);''')
doc.add_paragraph('')
doc.add_paragraph('Daftar UI primitives (21 komponen):', style='List Bullet')
for ui in ['alert-dialog', 'avatar', 'badge', 'breadcrumb', 'button', 'card', 'dialog', 'dropdown-menu', 'input', 'label', 'safe-image', 'select', 'separator', 'sheet', 'sidebar', 'skeleton', 'switch', 'table', 'tabs', 'textarea', 'tooltip']:
    doc.add_paragraph(ui, style='List Bullet')
doc.add_paragraph('')

doc.add_heading('Flat Components (Layout)', level=3)
add_path('src/components/ (14 komponen)')
doc.add_paragraph('Komponen yang dipakai di halaman publik (landing, paket, dll.):')
for comp in ['Navbar', 'Footer', 'HeroSection', 'AboutSection', 'ServicesSection', 'PackagesSection', 'GallerySection', 'TestimonialsSection', 'ContactSection', 'PackageDetailClient', 'RecommendationFinder', 'CountdownTimer', 'NotificationToast', 'ThemeProvider']:
    doc.add_paragraph(comp, style='List Bullet')
doc.add_paragraph('')

doc.add_heading('Panel Components', level=3)
add_path('src/components/panel/ (14 komponen)')
doc.add_paragraph('Komponen yang dipakai di dashboard admin & manager:')
for comp in ['AppSidebar', 'PanelHeader', 'PageHeader', 'StatCard', 'FormCard', 'SaveButton', 'ImageUpload', 'LocaleTabs', 'TranslateButton', 'ThemeToggle', 'ConfirmDialog', 'AccountClient', 'TopsisAnalysisClient', 'NotificationToast']:
    doc.add_paragraph(comp, style='List Bullet')
doc.add_paragraph('')

# ── 2.5 Multi-language ─────────────────────────────────────────────
doc.add_heading('2.5 Multi-language (en/id/ms)', level=2)
doc.add_paragraph(
    'Menggunakan LocaleContext (React Context) + file src/lib/translations.ts '
    '(2735 baris dictionary). JSON field di database menyimpan teks dalam 3 bahasa. '
    'LocaleTabs component untuk switch tab bahasa di form admin.'
)
add_path('src/lib/translations.ts')
doc.add_paragraph('')
add_path('src/lib/LocaleContext.tsx')
add_code('''const LocaleContext = createContext<{
  locale: string;
  setLocale: (locale: string) => void;
  t: (key: string) => string;
}>({ locale: "id", setLocale: () => {}, t: () => "" });''')
doc.add_paragraph('')

# ── 2.6 State Management ──────────────────────────────────────────
doc.add_heading('2.6 State Management Pattern', level=2)
doc.add_paragraph(
    'Tidak ada Redux/Zustand. Semua state dikelola via:',
    style='List Bullet'
)
doc.add_paragraph('React Context — ThemeProvider (dark/light), LocaleContext (bahasa)', style='List Bullet')
doc.add_paragraph('Server Actions — data fetching dan mutation langsung dari komponen client', style='List Bullet')
doc.add_paragraph('URL params — search params untuk filter, pagination', style='List Bullet')
doc.add_paragraph('useActionState — untuk form submission dengan pending state', style='List Bullet')
doc.add_paragraph('')

# ── 2.7 Export & Upload ────────────────────────────────────────────
doc.add_heading('2.7 Export & File Upload', level=2)
doc.add_paragraph('Export laporan ke Excel menggunakan exceljs + file-saver:', style='List Bullet')
add_path('src/lib/utils/export.ts')
doc.add_paragraph('')
doc.add_paragraph('Upload gambar ke Vercel Blob (dengan fallback local):', style='List Bullet')
add_path('src/lib/actions/upload.ts')
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════
# BAGIAN 3 — ALUR BELAJAR
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Bagian 3: Alur Belajar yang Disarankan', level=1)
doc.add_paragraph('')
learning_path = [
    ('1', 'Pahami Prisma Schema', 'Baca prisma/schema.prisma. Pahami model, relasi, enum.'),
    ('2', 'Baca Prisma Singleton', 'src/lib/prisma.ts — bagaimana koneksi DB dikelola.'),
    ('3', 'CRUD Sederhana', 'Baca 1 server action CRUD, misal bank-accounts.ts atau vehicle-types.ts (paling sederhana).'),
    ('4', 'API Route', 'Baca login route — pahami Request/Response, cookie, JWT.'),
    ('5', 'Auth System', 'src/lib/auth.ts — encrypt, decrypt, getSession.'),
    ('6', 'CRUD Kompleks', 'packages.ts — CRUD dengan relasi (priceTiers), multi-language JSON fields.'),
    ('7', 'TOPSIS Engine (Full)', 'src/lib/topsis.ts — pahami 7 langkah algoritma.'),
    ('8', 'TOPSIS Server Action', 'Baca src/lib/actions/topsis.ts — integrasi TOPSIS dengan DB.'),
    ('9', 'Root Layout & globals.css', 'Pahami font, provider, design tokens, utility classes.'),
    ('10', 'UI Primitives', 'Baca button.tsx atau badge.tsx — pahami cva, tailwind-merge.'),
    ('11', 'Flat Components', 'Baca Navbar.tsx atau HeroSection.tsx — client/server component.'),
    ('12', 'Panel Components', 'Baca AppSidebar.tsx, StatCard.tsx — komponen admin.'),
    ('13', 'Multi-language', 'LocaleContext, translations.ts, LocaleTabs — i18n tanpa library.'),
    ('14', 'Route Structure', 'Telusuri folder src/app — pahami App Router, route groups, layouts.'),
]

for num, title_text, desc in learning_path:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {title_text}')
    run.bold = True
    p.add_run(f' — {desc}')
doc.add_paragraph('')

# ── Simpan ─────────────────────────────────────────────────────────
output_path = '/mnt/d/job-besak/titan-travel/titan-travel-tutorial.docx'
doc.save(output_path)
print(f'✅ DOCX saved: {output_path}')
print(f'   Size: {os.path.getsize(output_path) / 1024:.1f} KB')
